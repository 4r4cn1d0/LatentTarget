"""Build a .docx of the MATS write-up (exec summary + main write-up + figures) for upload to Google Docs."""
import os, re
from docx import Document
from docx.shared import Inches, Pt
ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
W = os.path.join(ROOT, "results", "writeup"); OUT = os.path.join(W, "doc", "MATS_WRITEUP.docx")
mat = open(os.path.join(W, "WRITEUP_MATERIALS.md"), encoding="utf-8").read()
summary = open(os.path.join(W, "doc", "exec_summary_final.md"), encoding="utf-8").read().strip()
doc = Document(); st = doc.styles["Normal"]; st.font.name = "Arial"; st.font.size = Pt(11)

def runs(par, text):
    for tok in re.split(r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+\*)", text):
        if not tok: continue
        if tok.startswith("**"): r = par.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`"): r = par.add_run(tok[1:-1]); r.font.name = "Courier New"
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2: r = par.add_run(tok[1:-1]); r.italic = True
        else: par.add_run(tok)

def para(text, style=None):
    p = doc.add_paragraph(style=style); runs(p, text); return p

def section(title):
    m = re.search(r"^## %s.*?$" % re.escape(title), mat, flags=re.M)
    if not m: return ""
    start = m.end(); nxt = re.search(r"^## ", mat[start:], flags=re.M); return mat[start:start + nxt.start()] if nxt else mat[start:]

def md_block(text):
    lines = text.split("\n"); i = 0
    while i < len(lines):
        l = lines[i]
        if l.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].startswith("|"): rows.append(lines[i]); i += 1
            cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows if not re.match(r"^\|[-| ]+\|$", r.strip())]
            if cells:
                ncol = len(cells[0]); t = doc.add_table(rows=len(cells), cols=ncol); t.style = "Table Grid"
                for a, row in enumerate(cells):
                    for b in range(ncol):
                        c = row[b] if b < len(row) else ""; cell = t.cell(a, b); cell.text = ""; runs(cell.paragraphs[0], c)
                        for r in cell.paragraphs[0].runs: r.font.size = Pt(8); r.bold = bool(r.bold) or a == 0
                doc.add_paragraph()
            continue
        if l.startswith("- "): para(l[2:], style="List Bullet")
        elif l.startswith("## "): doc.add_heading(l[3:], level=2)
        elif l.startswith("# "): pass
        elif l.strip(): para(l)
        i += 1

def fig(name, caption):
    p = os.path.join(W, name)
    if os.path.exists(p):
        doc.add_picture(p, width=Inches(6.3)); c = doc.add_paragraph(); r = c.add_run(caption); r.italic = True; r.font.size = Pt(9)

doc.add_heading("LatentTarget: do LLMs learn which persuasion frame a hidden partner responds to?", level=0)
para("Aayush (Rishi) Ghosh · MATS 12.0 application · repository: github.com/4r4cn1d0/LatentTarget · every number below is generated from committed artifacts by scripts/make_writeup_materials.py and names its source.")
doc.add_heading("Executive summary", level=1)
for p in summary.split("\n\n"):
    if p.strip(): para(p.strip())
fig("fig_w11_whole_story.png", "Figure 1. Match rate by round with the model's own history. Left to right: Qwen3.8-27B under the V4 prompt (learns); the same model under a reworded prompt (Arm P1, still learns); Gemma-4-31B under the identical design (Arm R1, flat); the same Qwen forced to state a probability per message before choosing (Arm E1, flat). Dashed: V4 no-history reference. Grey band: held-out wording.")
doc.add_heading("Main write-up", level=1)
doc.add_heading("1. Question and hypotheses", level=2)
para("Does a model, told only to get Option A chosen and given a binary outcome each round, learn which persuasion frame (fairness, risk, expertise) a hidden partner responds to (H1), does that depend on its own history rather than any history (H2), does it generalise to unseen wording (H3), and does it revise after a silent change of partner (H4)? Behind these: is any learning a separable model of the partner, or a default preference nudged by feedback?")
doc.add_heading("2. Design (V4, frozen before the run)", level=2)
para("Each of 20 rounds shows a scenario with Option A and Option B and three unlabelled candidate messages, one per frame, drawn from a registered bank (90 templates, 45 held-out paraphrases). The model answers 1, 2, or 3. A simulated partner chooses A with probability 0.72 if the candidate's registered frame matches its hidden type and 0.38 otherwise. In swap episodes the type changes silently after round 10. Rounds 16 to 20 use held-out wording. Conditions: full history, no history, shuffled history (another episode's transcript), random partner, swap. 60 episodes per stable condition and 120 swap episodes: 360 episodes, 7,200 choices per run. Metrics: match rate; learning gain = late held-out match minus early match; difference-in-differences against no-history; after the swap, new-frame gain, old-frame drop, and late new-over-old; one-sided sign-flip randomization tests over episodes and episode bootstrap CIs. Thresholds, seeds, and the bank hash were frozen in a spec before any real run and each arm was analysed once.")
doc.add_heading("3. Results", level=2); md_block(section("Exec-summary skeleton"))
fig("fig_w1_v4_learning_by_condition.png", "Figure 2. V4 (Qwen3.8-27B): match rate by round for the four stable conditions.")
fig("fig_w6_v4_learning_by_target.png", "Figure 3. V4 per hidden target: learning is largest where the default frame is weakest.")
fig("fig_w2_v4_swap_by_transition.png", "Figure 4. V4 silent swap: new-frame gain and old-frame drop by transition; adaptation only into the default frame.")
fig("fig_w10_history_sensitivity_qwen_vs_gemma.png", "Figure 5. Does the choice depend on history at all? Win-stay/lose-shift and cross-condition agreement, Qwen vs Gemma.")
fig("fig_w9_elicited_belief_vs_choice.png", "Figure 6. Arm E1: stated belief vs choice by round and after the swap; the two lines coincide.")
fig("fig_w3_default_frame_priors.png", "Figure 7. No-history frame shares: the expertise default across banks and models.")
fig("fig_w5_first_crossing_bias.png", "Figure 8. Why a first-crossing 'probe leads behaviour' metric was dropped: a chance-level probe appears to lead.")
doc.add_heading("4. Gate ledger: every design and what stopped it", level=2); md_block(section("Gate ledger"))
doc.add_heading("5. Randomly selected examples (seeded, not chosen)", level=2)
para("The model sees the three candidates without frame labels; labels shown are the registered ground truth. → marks the model's choice.")
doc.add_heading("V4 (Qwen3.8-27B), seed 0", level=3); md_block(section("Randomly selected V4 examples"))
doc.add_heading("Arm R1 (Gemma-4-31B), seed 1", level=3); md_block(section("Randomly selected examples — Arm R1"))
doc.add_heading("Arm E1 (Qwen3.8-27B, elicited), seed 2", level=3); md_block(section("Randomly selected examples — Arm E1"))
doc.add_heading("6. Limitations", level=2)
para("One positive model out of two. The partner is a simulator with a fixed rule, not a human or an LLM. Persuasion is a choice among pre-written messages, not generated text. The expertise default confounds revision; four redesigns (V5 to V8) to remove it stopped at their own preregistered gates. Arm E1 cannot separate the output-format change from the model seeing its own past predictions. Arm P1 failed its validity gate (0.898) because the reworded prompt produced truncated reasoning preambles replaced by random fallbacks, which dilutes rather than inflates learning. The message bank was labelled by two blind machine judges; a 45-template blind hand-label sheet exists. No activation-level evidence. N = 60 episodes per stable condition.")
doc.add_heading("7. What I would do next", level=2)
para("A two-frame design with a distractor to remove the default; a larger N; an elicited arm whose history hides the model's own predictions, to separate format from anchoring; an LLM target with a persona; and only if a behavioural effect survives model and format changes, activation probes for the partner's type.")
doc.add_heading("Appendix A. Sourced numbers", level=2); md_block(section("Numbers sheet"))
doc.add_heading("Appendix B. Reproduction", level=2)
para("Specs: docs/behavioral_checkpoint_v4.json, docs/v4_replication_gemma4.json, docs/v4_elicited_qwen38.json, docs/v4_paraphrase_qwen38.json. Declarations and outcomes: docs/V4_REPLICATION_DECLARATION.md. Runner: scripts/run_controlled_open_weight.py; analyzers: scripts/analyze_controlled_v4.py, scripts/analyze_elicited_choices.py, scripts/analyze_elicited_beliefs.py. Log: docs/WORK_LOG.md. Tests: 793.")
doc.save(OUT); print("wrote", OUT, os.path.getsize(OUT) // 1024, "KB")

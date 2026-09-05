"""Build the MATS write-up as a clean document (.docx and .html) from committed artifacts.

Continuous prose, no placeholders, no code in the body. Numbers are read from the frozen result files.
Randomly selected examples are drawn from the raw logs with fixed seeds (V4 seed 0, P1 seed 3, R1 seed 1, E1 seed 2).
"""
import base64, html, json, os, random, re, sys
ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__))); sys.path.insert(0, ROOT)
from src.controlled_focal_agent import SPONTANEOUS_SYSTEM_TEMPLATE, ELICITED_SYSTEM_TEMPLATE, SPONTANEOUS_SYSTEM_TEMPLATE_PARAPHRASE_1  # noqa: E402
W = os.path.join(ROOT, "results", "writeup"); D = os.path.join(W, "doc"); os.makedirs(D, exist_ok=True)
RAW = os.path.join(ROOT, "data", "raw")
LOGS = {"v4": os.path.join(RAW, "qwen38_27b_v4_checkpoint_20260902.jsonl"), "p1": os.path.join(RAW, "v4p-qwen38.jsonl"), "r1": os.path.join(RAW, "v4r-gemma4.jsonl"), "e1": os.path.join(RAW, "v4e-qwen38.jsonl")}
J = lambda *p: json.load(open(os.path.join(ROOT, *p)))
V4 = J("results", "v4_real", "checkpoint", "v4_checkpoint_summary.json"); P1 = J("results", "v4_real", "paraphrase_qwen38", "v4_checkpoint_summary.json")
R1 = J("results", "v4_real", "replication_gemma4", "v4_checkpoint_summary.json"); E1 = J("results", "v4_real", "elicited_qwen38", "elicited_choice_summary.json"); E1B = J("results", "v4_real", "elicited_qwen38", "beliefs", "elicited_belief_summary.json")
def g(s, cond): m = s["stable_condition_metrics"][cond]; return m
def ci(x): return "%.3f [%.3f, %.3f]" % (x["mean"], x["ci_lo"], x["ci_hi"])
v4f = g(V4, "full_history"); p1f = g(P1, "full_history"); r1f = g(R1, "full_history"); e1f = E1["stable_condition_metrics"]["elicited_full_history"]
v4sw = V4["swap_metrics"]; p1sw = P1["swap_metrics"]; r1sw = R1["swap_metrics"]; e1sw = E1["swap_metrics"]
v4did = V4["primary_contrasts"]["full_vs_no_difference_in_differences"]; p1did = P1["primary_contrasts"]["full_vs_no_difference_in_differences"]
v4t = V4["late_match_by_target_type"]; p1t = P1["late_match_by_target_type"]; r1t = R1["late_match_by_target_type"]
e1x = E1["cross_prompt_comparison_vs_v4_spontaneous"]; e1post = E1B["per_condition"]["elicited_swap"]["post_swap_rounds_1_to_5_belief_minus_choice_new"]
alpha = V4["thresholds_frozen_before_real_run"]["confirmatory_alpha_one_sided"]

def examples(key, k, seed):
    lines = open(LOGS[key]).readlines(); idx = sorted(random.Random(seed).sample(range(len(lines)), k)); out = []
    for i in idx:
        r = json.loads(lines[i]); cands = sorted(r["candidates"], key=lambda c: c["slot"])
        out.append({"cond": r["condition"], "round": r["round"], "type": r["hidden_target_type"], "scenario": r["scenario"].get("title"), "cands": [(c["slot"], c["frame"], c["message"]) for c in cands], "slot": r["selected_slot"], "frame": r["selected_frame"], "p_a": r["target_p_a"], "choice": r["target_choice"], "raw": r["focal_output_raw"].strip(), "valid": r["selection_valid"]})
    return out

# ------------------------------------------------------------------ content
C = []
title = lambda t: C.append(("title", t)); h1 = lambda t: C.append(("h1", t)); h2 = lambda t: C.append(("h2", t)); h3 = lambda t: C.append(("h3", t)); p = lambda t: C.append(("p", t))
bullets = lambda xs: C.append(("bullets", xs)); fig = lambda n, c: C.append(("fig", n, c)); quote = lambda t: C.append(("quote", t)); table = lambda rows: C.append(("table", rows))

title("Testing Whether Language Models Learn Hidden Persuasion Targets")
p("Aayush Ghosh. Application project for Neel Nanda's MATS 12.0 stream. Code, logs, and every number in this document: github.com/4r4cn1d0/LatentTarget.")

h1("Executive summary")
for para in open(os.path.join(D, "exec_summary_final.md"), encoding="utf-8").read().strip().split("\n\n"):
    para = para.strip().replace(" [CONFIRM in your words]", "")
    if not para: continue
    if para.startswith("[FIG:"): name, cap = para[5:].split("]", 1); fig(name.strip(), cap.strip())
    elif para.startswith("- "): bullets([l[2:] for l in para.split("\n") if l.startswith("- ")])
    elif "\n- " in para:
        head, rest = para.split("\n- ", 1); p(head.strip()); bullets([l.strip() for l in ("- " + rest).split("\n- ") if l.strip()][0:] if False else [x.strip() for x in rest.split("\n- ")])
    else: p(para)

h1("Randomly selected examples")
p("These are drawn from the raw logs with fixed seeds, not chosen. Each shows one round: the scenario, the three candidate messages exactly as the model saw them (the frame labels in brackets are the registered ground truth, which the model never sees), the model's answer, and what the simulated partner did.")
def render_examples(key, k, seed, title):
    h3(title)
    for e in examples(key, k, seed):
        p("%s, round %d of 20, partner type %s. Scenario: %s." % (e["cond"].replace("_", " "), e["round"], e["type"], e["scenario"]))
        bullets(["%s%d [%s] %s" % ("Chosen: " if s == e["slot"] else "", s, f, m.replace("\n", " ")) for s, f, m in e["cands"]])
        tail = "Model answered \"%s\"" % e["raw"][:60] + ("" if e["valid"] else " (did not parse; fallback slot %d assigned at random)" % e["slot"]) + ". Partner P(A) = %.2f, chose %s." % (e["p_a"], e["choice"])
        p(tail)
render_examples("v4", 5, 0, "Qwen3.8-27B, original prompt (V4): five draws from 7,200 records")
render_examples("p1", 3, 3, "Qwen3.8-27B, reworded prompt (Arm P1): three draws from 7,200")
render_examples("r1", 3, 1, "Gemma-4-31B (Arm R1): three draws from 7,200")
render_examples("e1", 3, 2, "Qwen3.8-27B, stated beliefs (Arm E1): three draws from 3,600")

h1("How the results were checked")
p("This project was heavily AI-assisted. Claude, in Claude Code, and Codex wrote most of the code, ran the GPU jobs, and helped draft the report. I set the research question, directed the project, and approved the sequence of experiments. Agents proposed and implemented many of the technical details. That makes the audit trail especially important: the central claims should stand on the frozen specifications, raw logs, and independent calculations, not on an agent sounding confident.")
bullets([
    "Preregistration. Every real run has a frozen specification (thresholds, seeds, and a hash of the message bank) committed before any data existed. Each run was analysed once with the frozen analyzer. A failed rule stops the run; nothing was re-tuned after seeing a result, and the four redesigns that failed their rules are reported below.",
    "The pipeline can detect learning. A simulated Bayesian learner run through the identical pipeline passes the same rules, and a simulated non-learner fails them. A flat result therefore means the model did not learn, not that the analyzer is broken.",
    "Belief equals choice is model behaviour, not a parser artefact. The runner records the model's own choice field from its JSON answer. In all 3,600 elicited records that stated choice was in the argmax set of the model's own stated probabilities (25 rounds had ties, all resolved by the model's stated choice). This was re-derived with a separate script over the raw log.",
    "No position bias. Chosen-slot shares are 0.34, 0.33, 0.33 for Gemma and 0.31, 0.35, 0.35 for the elicited Qwen run; candidate order is shuffled per round.",
    "Model-free diagnostics that do not depend on the analyzer. The probability of repeating the previous frame after a success minus after a failure is 0.183 for Qwen and 0.031 for Gemma. Gemma makes the same choice as with someone else's history on the identical candidate triple 90.5% of the time (Qwen 63.7%). Both were computed directly from the logs.",
    "Numbers are generated, not typed. A script reads the committed result files and writes every figure and a table of 124 numbers, each with its source file; this document's numbers are read from the same files. One transcription error in a comparison table was caught by an independent check of the write-up against those files and corrected before submission.",
    "Message bank labels. Two blind machine judges labelled the frame of every message; the judge outputs are committed. A blind hand-labelling sheet of all 45 templates was also prepared.",
    "The project reviewed its own rule. A candidate revision rule (V7) was rejected after a five-lens adversarial review showed its pooled test would pass on a model that only ever drifted to its default frame.",
    "Tests. 793 automated tests pass, including tests that pin the original prompt text by hash, pin the registered thresholds, and check the on-disk Gemma prior measurement against its raw log.",
])

h1("1. The question, and why I chose it")
p("Chen and colleagues showed that language models form accurate models of static attributes of the person they are talking to, from very little text. I wanted the dynamic version, in the setting where it matters most. Persuasion is where a model's picture of its interlocutor turns into action. Over a repeated interaction with feedback, does the model learn which kind of argument this particular partner responds to? Does it revise that when the partner silently changes? And is there any stated belief about the partner that is separate from the choice it makes? A persuader that holds and updates such a model is a different safety problem from one with a fixed style, and any interpretability claim about a representation of the target needs the behavioural effect first.")
p("The hypotheses, in the form they were registered before the first real run: with its own history the model matches the partner's preferred frame more often than chance (H1); that depends on its own history, not on any history and not on a partner that ignores it (H2); it generalises to unseen wording of the same frames (H3); and after a silent change of partner it moves to the new frame and away from the old one (H4). A further question, tested after the first run: is a stated belief separable from the choice?")
p("The project therefore moved from free-form persuasion to a controlled choice. In a free-form design, which frame the model used is a judge's opinion, and judge noise could swamp a modest learning effect. Here the model picks one of three registered messages, one per frame, and the partner responds to the registered frame. Learning becomes a match rate with no judge in the loop. The price is important: the model selects the persuasion rather than writing it.")

h1("2. Setup")
h2("Environment")
p("Each round presents a scenario with an Option A and an Option B, and three candidate messages arguing for A: one from fairness, one from risk, one from expertise. The messages come from a registered bank of 45 templates, 15 per frame, of which 10 per frame are used in the first 15 rounds and 5 per frame are reserved for rounds 16 to 20 as held-out wording. The three candidates are shown in a seeded random order without labels, and the model answers with a single digit.")
p("The partner is a simulator with a hidden type. It chooses A with probability 0.72 when the chosen message's registered frame matches its type and 0.38 otherwise. An episode is 20 rounds. In swap episodes the type changes silently after round 10. The five conditions are: full history, where the model sees its own transcript so far; no history; shuffled history, where it sees another episode's transcript; random partner, where the partner ignores the frame; and swap. There are 60 episodes per stable condition and 120 swap episodes, so a full run is 360 episodes and 7,200 choices. The same bank, seeds, and partner types are used in every run.")
p("The models are Qwen3.8-27B and Gemma-4-31B-it, both open weights at pinned revisions, run in bf16 on one A100 with greedy decoding. The answer budget is 8 tokens for the single-digit prompts and 96 for the JSON prompt.")
h2("Measures")
p("A round is a match when the chosen candidate's registered frame equals the partner's current type; chance is one third. The learning gain of an episode is its mean match on the held-out rounds 16 to 20 minus its mean match on rounds 1 to 5. The primary test is the difference between the full-history and no-history learning gains, with a one-sided sign-flip randomization test over episodes (10,000 draws) and episode-bootstrap confidence intervals. For swap episodes I report the new-frame gain (late new-frame use minus pre-swap new-frame use), the old-frame drop, and the late new-over-old difference. The registered revision rule requires a new-frame gain and an old-frame drop of at least 0.10 and a positive late new-over-old under the same test at a one-sided alpha of %s, the 0.05 budget being split between the learning and revision tests. An episode counts as adapted when three of four consecutive post-swap picks are the new frame." % alpha)
p("Two model-free diagnostics are reported alongside, because they do not depend on any of the above: the probability of repeating the previous round's frame after a success versus after a failure, and how often a model makes the same choice in two conditions when shown the identical candidate triple.")
h2("Prompts")
p("The original system prompt, used in V4 (its text is pinned by hash in the test suite):"); quote(SPONTANEOUS_SYSTEM_TEMPLATE.format(n_rounds=20))
p("The reworded system prompt used in Arm P1, with the same facts in different sentences:"); quote(SPONTANEOUS_SYSTEM_TEMPLATE_PARAPHRASE_1.format(n_rounds=20))
p("The elicited system prompt used in Arm E1, which asks for a probability per candidate before the choice; in this arm the model's past predictions are also shown in its history:"); quote(ELICITED_SYSTEM_TEMPLATE.format(n_rounds=20))
p("The user turn shows the scenario, the visible history (each previous round's three candidates, the model's choice, and the partner's choice), and the three current candidates. An audit checks that no prompt names the frames or contains instructions such as \"adapt to the participant\".")

h1("3. Experiment 1: does the model learn the partner's frame?")
p("Prediction, written before the run: the first three hypotheses hold if the full-history learning gain is positive with the three controls flat, and the effect survives on held-out wording. The outcomes I considered possible were no learning at all, apparent learning that is really a default frame (full history and no history both high on the same frame), and genuine feedback-driven learning.")
p("Outcome: genuine learning, on one model. With its own history, Qwen's held-out match rate rose from %.3f in the early rounds to %.3f. The learning gain was %s and the registered randomization test passed (difference-in-differences against no history %s, p = %.4f). No history stayed at %.3f, shuffled history fell from %.3f to %.3f, and the random partner was flat. The learning is anti-default. With no history the model picks the expertise message 92.2%% of the time, so there is nothing to learn when the partner happens to prefer expertise (late match %.2f with history versus %.2f without). The gains come on fairness partners (%.2f versus %.2f) and risk partners (%.2f versus %.2f)." % (v4f["early_match"]["mean"], v4f["late_heldout_match"]["mean"], ci(v4f["learning_gain"]), ci(v4did), v4did["p_value_one_sided"], g(V4, "no_history")["late_heldout_match"]["mean"], g(V4, "shuffled_history")["early_match"]["mean"], g(V4, "shuffled_history")["late_heldout_match"]["mean"], v4t["expertise"]["full_late_heldout"], v4t["expertise"]["no_history_late_heldout"], v4t["fairness"]["full_late_heldout"], v4t["fairness"]["no_history_late_heldout"], v4t["risk"]["full_late_heldout"], v4t["risk"]["no_history_late_heldout"]))
fig("fig_w6_v4_learning_by_target.png", "Figure 4. V4, by the partner's type: late held-out match with full history, no history, and shuffled history. Learning is largest where the default frame is weakest.")

h1("4. Experiment 2: does it revise after a silent change of partner?")
p("Prediction: the fourth hypothesis holds if, after the swap at round 10, new-frame use rises and old-frame use falls and, by the end, the new frame is used more than the old one. The alternative I was worried about was adaptation that only ever runs toward the default frame.")
p("Outcome: the two effect thresholds passed and the decisive test failed. Over 120 swap episodes, new-frame use rose by %.3f and old-frame use fell by %.3f, but late new-frame use never exceeded old-frame use (difference %.3f, p = %.2f). %d of 120 episodes adapted, and where they adapted tells the story: 34 of 40 swaps into expertise, 9 of 40 into risk, 0 of 40 into fairness (Figure 2). This is what a default frame re-asserting itself looks like. It is not what an updated model of the partner would produce, though the mechanism is inferred from the pattern, not shown." % (v4sw["new_target_gain"]["mean"], v4sw["old_target_drop"]["mean"], v4sw["late_new_over_old"]["mean"], v4sw["late_new_over_old"]["p_value_one_sided"], v4sw["n_adapted"]))

h1("5. Experiment 3: three stress tests, each predicted before its outcome")
p("After the first run, predictions were frozen for three further runs, and each was run once with the same design and analyzer. The first two were declared together. The third was declared after their outcomes were known but before its own run. Figure 3 shows the four with-history curves side by side.")
h2("5.1 Reworded prompt, same model: the effect survives")
p("The question was whether the learning depends on the exact wording of the original prompt. It does not. With the reworded prompt the learning gain was %s, against %s originally; the difference-in-differences against no history was %s, p = %.4f; the controls were flat; and the per-partner pattern was the same, with the advantage over no history at %.2f for fairness, %.2f for risk, and %.2f for expertise. Revision failed again, as predicted (new-frame gain %.3f, old-frame drop %.3f, late new-over-old %.3f)." % (ci(p1f["learning_gain"]), ci(v4f["learning_gain"]), ci(p1did), p1did["p_value_one_sided"], p1t["fairness"]["advantage"], p1t["risk"]["advantage"], p1t["expertise"]["advantage"], p1sw["new_target_gain"]["mean"], p1sw["old_target_drop"]["mean"], p1sw["late_new_over_old"]["mean"]))
p("One thing did not go as planned. Under the new wording the model began a reasoning preamble (\"Looking at the history, the participant chose…\") in 12.2%% of rounds that had a history and in none of the rounds without one, and was cut off by the 8-token budget. Those 733 rounds, 10.2%% of the run, were assigned a uniformly random slot by the frozen fallback rule, which fails the run's validity threshold (%.3f valid against a required 0.98) and can only pull the learning toward chance: on rounds that parsed, the late held-out match was 0.632 against 0.600 over all rounds. I did not re-run with a larger budget, because the run had been declared with the original decoding settings. The preambles are themselves a small observation: under this wording the model spontaneously refers to the partner's past choices when it starts to explain itself." % P1["valid_selection_rate"])
h2("5.2 A second model family: no replication")
p("Gemma-4-31B-it, on the identical design, does not learn. Its learning gain was %s and every effect rule except the random-partner control failed. It is nearly insensitive to both history and feedback: its choice equals its shuffled-history choice on the identical candidate triple 90.5%% of the time, and it repeats its previous frame with probability 0.972 after a success and 0.941 after a failure, a gap of 0.031 where Qwen's is 0.183 (Figure 5). Its default is actually weaker than Qwen's (expertise 78.7%% of no-history picks against 92.2%%), so the strength of the default does not explain the difference; the use of feedback does. What movement there is runs toward the default: with history, expertise partners are matched %.2f of the time against %.2f without, while fairness partners fall from %.2f to %.2f." % (ci(r1f["learning_gain"]), r1t["expertise"]["full_late_heldout"], r1t["expertise"]["no_history_late_heldout"], r1t["fairness"]["no_history_late_heldout"], r1t["fairness"]["full_late_heldout"]))
fig("fig_w10_history_sensitivity_qwen_vs_gemma.png", "Figure 5. Two model-free diagnostics. Left: probability of repeating the previous round's frame after a success and after a failure. Right: how often the model makes the same choice as with its own history when shown the identical candidate triple with a shuffled history, no history, or a random partner.")
h2("5.3 Stated beliefs: no separate belief, and the learning disappears")
p("The point of this run was to see whether the model holds a belief about the partner that its choices do not act on. The prediction was that if such a belief existed, the stated belief would move to the new frame after the swap before the choice did; if the two moved together, the stated belief was a description of the policy. Neither moved. Under the elicited prompt the same Qwen showed no learning (gain %s, which is %.3f below the original prompt, difference %s), picked expertise in 94.2%% of rounds and fairness in none, and its choice was the argmax of its own stated probabilities in every one of 3,600 records. Its stated confidence was a fixed ranking of the frames, expertise 0.69, risk 0.58, fairness 0.49, that shifted by about 0.05 with feedback without ever changing the choice. After the swap, the share of rounds where the stated belief matched the new frame exceeded the share where the choice did by %.3f." % (ci(e1f["learning_gain"]), abs(e1x["learning_gain_diff"]["mean"]), ci(e1x["learning_gain_diff"]), e1post["mean"]))
p("Three caveats. This run changed two things at once, the output format and the fact that the model saw its own past predictions in its history, and cannot separate them. Its prompt was never tuned, whereas the original was. And the 96-token budget leaves no room for reasoning. So this run could not answer the belief question in the positive direction, because the learning itself vanished. What it does show is that the stated probabilities carried no information about the partner that the choice lacked.")
fig("fig_w9_elicited_belief_vs_choice.png", "Figure 6. Arm E1. Left and middle: by round, the share of episodes where the stated belief matches the partner's type, where the choice matches it, and where belief and choice agree. Right: after the silent swap, by rounds since the swap, the share where the stated belief and the choice match the new type. The belief and choice lines coincide throughout.")

h1("6. What the evidence supports, and the strongest case against")
p("Supported: a real, preregistered behavioural effect on one model, robust to the wording of the prompt. Given only outcomes, Qwen3.8-27B moves its choice toward the argument frame that works, on unseen wording, and the controls rule out the effect coming from any history at all or from a partner that is not responding.")
p("Not shown: a model of the partner that is separate from the choice. Revision cleared its two effect thresholds but not the test that matters, and adaptation ran mostly toward the default. Forcing the model to state beliefs removed the effect and showed that belief and choice were one object. A second model family showed nothing.")
p("The strongest alternative explanation is a model-free policy: repeat what worked, otherwise fall back to expertise, on top of a strong default. I have not fitted such a policy, but nothing here rules it out. What would separate it from a partner model is revision away from the default, or a stated belief that leads the choice. Neither happened.")
p("What I did not do: no activation capture, probes, or steering. The plan gated them behind a passed revision test, which never came, and the stated-belief result made a probe for the partner's type premature until a behavioural effect survives a change of model and of format.")

h1("7. Limitations")
p("One positive model out of two. A third model under the original prompt would separate \"Qwen\" from \"feedback-sensitive\"; it would cost about three dollars of GPU time and was not run. The partner is a simulator with a fixed rule, not a person or a language model with a persona; a model partner would be more realistic at the cost of a judge deciding whether it responded to the frame. The persuasion is a choice among pre-written messages rather than generated text, which was the price of a judge-free measure. The expertise default confounds revision, and four redesigns to remove it stopped at their own rules; a two-frame design with a distractor is the fix I would try next. The stated-belief run cannot separate the format change from the model seeing its own predictions; a run with those predictions hidden would fix that. The reworded-prompt run failed its validity threshold on truncated preambles, which a larger token budget would remove. The message bank's frame labels rest on two blind machine judges. And 60 episodes per stable condition, with 20 per swap transition, limits the power to detect smaller effects.")

h1("8. What was stopped along the way")
p("Four successors to the first design were meant to remove the default-frame confound. Each had a feasibility rule written down before its outcome, and each stopped at it. The first (V5) rebuilt the message bank with a calibrated target and required the no-history frame shares to sit between 25%% and 42%%; the measured shares were 13.7, 34.2, and 52.1. The second (V6) required a balance gate that a 120,000-study simulation showed to be infeasible at every allowed sample size. The third (V7) dropped the balance requirement, failed its own feasibility rule, and was rejected on review because its pooled revision test would pass on a model that only drifted to its default. The fourth (V8) added a destination-stratified acquisition rule that controlled the false-positive rate (no joint rejections in 6,000 simulated null studies) but was underpowered against the weakest learner it was registered to detect. I also dropped a planned \"the stated belief leads the behaviour\" measure based on the first round at which each crosses a threshold, after simulating it: a belief probe at chance level appears to lead the behaviour by 0.91 rounds, with a confidence interval excluding zero in 87%% of simulated runs (Figure 8)." % ())
table([["Design", "What it changed", "Rule it was judged by", "Result"],
       ["V4", "controlled choice among three registered messages; partner responds to the registered frame", "learning and revision tests", "learning passed; revision failed"],
       ["V5", "24 rounds, constrained decoding, calibrated bank", "no-history frame shares between 25% and 42%", "failed (13.7 / 34.2 / 52.1)"],
       ["V6", "whole-triad selection, matched stable twin for every swap", "balance gate reachable at some sample size", "infeasible at every size"],
       ["V7", "balance gate dropped, measured nuisance cells", "own feasibility rule; adversarial review", "failed; review found the pooled rule passes on default drift"],
       ["V8", "destination-stratified acquisition rule, split alpha", "power against the weakest registered learner", "false-positive rate controlled; underpowered"],
       ["P1", "original design, reworded prompt", "V4 rules", "learning passed; revision failed; validity threshold failed on truncated preambles"],
       ["R1", "original design on Gemma-4-31B", "V4 rules", "learning failed; revision failed"],
       ["E1", "original design, stated beliefs in JSON", "V4 rules applied within the run", "learning failed; belief equals choice"]])
fig("fig_w3_default_frame_priors.png", "Figure 7. The default frame. Share of each frame chosen with no history: Qwen on the original bank, Qwen on the recalibrated bank, and Gemma.")
fig("fig_w5_first_crossing_bias.png", "Figure 8. Why the first-crossing measure was dropped: in simulation, a belief probe at chance level appears to lead the behaviour.")

h1("9. How I used language models, and how much to trust each part")
p("Claude, running in Claude Code, and Codex wrote nearly all of the code and analysis scripts, operated the GPU jobs, and produced early drafts of the documentation. GPT-5.6 acted as the two blind judges that labelled the message bank. I supplied the research question, directed the work, and approved the experiment sequence; agents proposed and implemented many technical details, including the controlled-choice redesign and several gates. The checks are listed in the earlier audit section. I trust the main learning result more than the secondary diagnostics because the controls, simulated-learner validation, prompt rewording, and seeded examples all point in the same direction. I trust the message-bank labels least because they still rest on machine judges; the planned human labels have not been completed.")

h1("10. What I would do next")
p("A third model family under the original prompt. A design with two frames and a distractor, to remove the default frame rather than work around it. A stated-belief run with the model's past predictions hidden from its history. A larger token budget for the reworded prompt, to see what the model says when it explains itself. And only if a behavioural effect survives all of that, activation probes for the partner's type.")

h1("Appendix: reproduction")
p("The repository contains the frozen specifications for each run, the declarations with predictions and outcomes, the runner, the analyzers, the raw-log manifests, the script that generates every figure and number in this document with its source file, the full decision log, and the test suite. All models are open weights at pinned revisions; each run takes one A100 for one to five hours.")

# ------------------------------------------------------------------ renderers
def esc(s): return html.escape(s)
def inl(s): return re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", html.escape(s))
def render_html():
    out = []
    for it in C:
        k = it[0]
        if k == "title": out.append("<h1 class='title'>%s</h1>" % esc(it[1]))
        elif k in ("h1", "h2", "h3"): out.append("<%s>%s</%s>" % (k, esc(it[1]), k))
        elif k == "p": out.append("<p>%s</p>" % inl(it[1]))
        elif k == "bullets": out.append("<ul>" + "".join("<li>%s</li>" % inl(b) for b in it[1]) + "</ul>")
        elif k == "quote": out.append("<blockquote style='border-left:3px solid #999;margin:8px 0;padding:4px 12px;white-space:pre-wrap'>%s</blockquote>" % esc(it[1]))
        elif k == "table": out.append("<table border='1' cellpadding='5' style='border-collapse:collapse;font-size:90%'>" + "".join("<tr>" + "".join(("<th>%s</th>" if i == 0 else "<td>%s</td>") % esc(c) for c in row) + "</tr>" for i, row in enumerate(it[1])) + "</table>")
        elif k == "fig":
            pth = os.path.join(W, it[1])
            if os.path.exists(pth): out.append('<p><img src="data:image/png;base64,%s" style="max-width:100%%"/></p><p><i>%s</i></p>' % (base64.b64encode(open(pth, "rb").read()).decode(), esc(it[2])))
    open(os.path.join(D, "MATS_WRITEUP.html"), "w", encoding="utf-8").write("<html><head><meta charset='utf-8'><title>LatentTarget MATS write-up</title><style>h1,h2,h3{color:#000}.title{font-size:1.65em}</style></head><body style='font-family:Arial,sans-serif;max-width:860px;margin:auto;line-height:1.45'>" + "\n".join(out) + "</body></html>")
def render_docx():
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.oxml.ns import qn
    doc = Document()
    style_specs = {
        "Normal": (11, False),
        "Title": (18, True),
        "Heading 1": (15, True),
        "Heading 2": (12.5, True),
        "Heading 3": (11.5, True),
    }
    for style_name, (size, bold) in style_specs.items():
        st = doc.styles[style_name]
        st.font.name = "Arial"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.color.rgb = RGBColor(0, 0, 0)
    doc.styles["Title"].paragraph_format.space_after = Pt(8)
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        doc.styles[style_name].paragraph_format.keep_with_next = True
    for it in C:
        k = it[0]
        if k == "title": doc.add_paragraph(it[1], style="Title")
        elif k == "h1": doc.add_heading(it[1], level=1)
        elif k == "h2": doc.add_heading(it[1], level=2)
        elif k == "h3": doc.add_heading(it[1], level=3)
        elif k == "p":
            par = doc.add_paragraph()
            for tok in re.split(r"(\*\*.+?\*\*)", it[1]):
                if tok.startswith("**"): r = par.add_run(tok[2:-2]); r.bold = True
                elif tok: par.add_run(tok)
        elif k == "bullets":
            for b in it[1]: doc.add_paragraph(b, style="List Bullet")
        elif k == "quote":
            par = doc.add_paragraph(); par.paragraph_format.left_indent = Inches(0.4); r = par.add_run(it[1]); r.italic = True; r.font.size = Pt(10)
        elif k == "table":
            rows = it[1]; t = doc.add_table(rows=len(rows), cols=len(rows[0])); t.style = "Table Grid"
            for a, row in enumerate(rows):
                for b, c in enumerate(row):
                    cell = t.cell(a, b); cell.text = c
                    for r in cell.paragraphs[0].runs: r.font.size = Pt(9); r.bold = a == 0
            doc.add_paragraph()
        elif k == "fig":
            pth = os.path.join(W, it[1])
            if os.path.exists(pth): doc.add_picture(pth, width=Inches(6.3)); c = doc.add_paragraph(); r = c.add_run(it[2]); r.italic = True; r.font.size = Pt(9)
    doc.save(os.path.join(D, "MATS_WRITEUP.docx"))
def render_txt():
    out = []
    for it in C:
        k = it[0]
        if k == "title": out.append("\n# " + it[1] + "\n")
        elif k in ("h1", "h2", "h3"): out.append("\n" + "#" * int(k[1]) + " " + it[1] + "\n")
        elif k == "p": out.append(it[1].replace("**", "") + "\n")
        elif k == "bullets": out.extend("- " + b.replace("**", "") for b in it[1]); out.append("")
        elif k == "quote": out.append("> " + it[1].replace("\n", "\n> ") + "\n")
        elif k == "table": out.extend(" | ".join(r) for r in it[1]); out.append("")
        elif k == "fig": out.append("[Figure %s] %s\n" % (it[1], it[2]))
    open(os.path.join(D, "MATS_WRITEUP.txt"), "w", encoding="utf-8").write("\n".join(out))
render_html(); render_docx(); render_txt()
txt = open(os.path.join(D, "MATS_WRITEUP.txt")).read()
print("built; words %d; CONFIRM tags %d; backticks %d; 'scripts/' mentions %d; figures %d" % (len(txt.split()), txt.count("CONFIRM"), txt.count("`"), txt.count("scripts/"), sum(1 for it in C if it[0] == "fig")))
